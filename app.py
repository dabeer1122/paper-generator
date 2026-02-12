import streamlit as st
import random
import qrcode
from docx import Document
from docx.shared import Inches
from io import BytesIO
import PyPDF2
import re
from datetime import datetime

st.title("📘 Personal Paper Generator System")

uploaded_file = st.file_uploader("Upload Word or PDF File", type=["pdf", "docx"])

def extract_text_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_mcqs(text):
    questions = re.split(r'\d+\.', text)
    mcqs = []
    for q in questions:
        if "A)" in q and "B)" in q and "C)" in q and "D)" in q:
            mcqs.append(q.strip())
    return mcqs

if uploaded_file is not None:
    if uploaded_file.type == "application/pdf":
        text = extract_text_from_pdf(uploaded_file)
    else:
        st.warning("Please upload PDF format for now.")
        st.stop()

    mcqs = extract_mcqs(text)

    if len(mcqs) < 15:
        st.error("Not enough MCQs found (minimum 15 required)")
    else:
        selected = random.sample(mcqs, 15)

        version = st.selectbox("Select Version", ["A", "B"])

        if st.button("Generate Paper"):

            doc = Document()
            doc.add_heading("BOARD OF INTERMEDIATE & SECONDARY EDUCATION", level=1)
            doc.add_paragraph("Health & Physical Education")
            doc.add_paragraph(f"Version {version}")

            paper_code = "PHE-" + datetime.now().strftime("%d%m%y%H%M%S")
            doc.add_paragraph(f"Paper Code: {paper_code}")

            qr = qrcode.make(paper_code)
            qr_bytes = BytesIO()
            qr.save(qr_bytes)
            qr_bytes.seek(0)
            doc.add_picture(qr_bytes, width=Inches(1.2))

            doc.add_paragraph("Name: _______________________")
            doc.add_paragraph("Roll No: _______________________")
            doc.add_paragraph("")

            for i, q in enumerate(selected, 1):
                doc.add_paragraph(f"{i}. {q}")

            doc.add_page_break()
            doc.add_heading("OMR SHEET", level=2)

            for i in range(1, 16):
                doc.add_paragraph(f"{i})  ○A   ○B   ○C   ○D")

            doc.add_page_break()
            doc.add_heading("Answer Key", level=2)

            for i, q in enumerate(selected, 1):
                ans = re.search(r'Answer[:\- ]+([A-D])', q)
                if ans:
                    doc.add_paragraph(f"{i}. {ans.group(1)}")

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="Download Paper",
                data=buffer,
                file_name="Generated_Paper.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
      )
